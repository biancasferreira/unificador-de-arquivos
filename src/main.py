# ==============================================================================
# Projeto: Unificador de Arquivos
# Criado por: Bianca Dos Santos
# Data de criação:20/02/2025
# Descrição: Este projeto unifica arquivos TXT, PDF e Word em um único arquivo.
# GitHub: https://github.com/biancasferreira/unificador-de-arquivos
# ==============================================================================

import os
import tkinter as tk
from tkinter import ttk, filedialog, scrolledtext, messagebox
from PyPDF2 import PdfMerger
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# Função para unificar os arquivos, considerando o tipo selecionado
def process_files(directory, remove_lines, insert_lines, add_counter, file_type):
    output_file = os.path.join(directory, "resultado")
    
    if file_type == "Texto":
        output_file += ".txt"
    elif file_type == "PDF":
        output_file += ".pdf"
    elif file_type == "Word":
        output_file += ".docx"
    
    if os.path.exists(output_file):
        os.remove(output_file)
        log_message(f"Arquivo de saída existente removido: {output_file}")

    if file_type == "Texto":
        files = [f for f in os.listdir(directory) if f.lower().endswith(".txt")]
    elif file_type == "PDF":
        files = [f for f in os.listdir(directory) if f.lower().endswith(".pdf")]
    elif file_type == "Word":
        files = [f for f in os.listdir(directory) if f.lower().endswith(".docx")]
    else:
        files = []

    if not files:
        log_message(f"Nenhum arquivo {file_type} encontrado.")
        return

    if file_type == "Texto":
        all_lines = []
        for filename in files:
            filepath = os.path.join(directory, filename)
            log_message(f"Processando: {filepath}")
            try:
                with open(filepath, "r", encoding="utf-8") as f:
                    lines = f.readlines()

                if remove_lines:
                    if lines[0].startswith("1;00;A"):
                        lines = lines[1:]
                    if lines and lines[-1].startswith("1;99;A"):
                        lines = lines[:-1]

                all_lines.extend(lines)
            except Exception as e:
                log_message(f"  -> Erro ao processar {filename}: {e}")

        if insert_lines:
            all_lines.insert(0, "1;00;A;02/01/2025;\n")

        if add_counter:
            count = len(all_lines) - (1 if insert_lines else 0)
            all_lines.append(f"1;99;A;{count};\n")

        with open(output_file, "w", encoding="utf-8") as out_f:
            out_f.writelines(all_lines)
        log_message("Processamento concluído!")

    elif file_type == "PDF":
        pdf_merger = PdfMerger()
        for filename in files:
            filepath = os.path.join(directory, filename)
            log_message(f"Processando PDF: {filepath}")
            try:
                pdf_merger.append(filepath)
            except Exception as e:
                log_message(f"  -> Erro ao processar PDF {filename}: {e}")

        pdf_merger.write(output_file)
        pdf_merger.close()
        log_message("Processamento concluído!")

    elif file_type == "Word":
        doc = Document()
        for filename in files:
            filepath = os.path.join(directory, filename)
            log_message(f"Processando Word: {filepath}")
            try:
                doc_to_add = Document(filepath)
                for para in doc_to_add.paragraphs:
                    doc.add_paragraph(para.text)
            except Exception as e:
                log_message(f"  -> Erro ao processar Word {filename}: {e}")

        doc.save(output_file)
        log_message("Processamento concluído!")

# Função de log
def log_message(message):
    text_log.insert(tk.END, message + "\n")
    text_log.see(tk.END)

# Função para selecionar diretório
def select_directory():
    directory = filedialog.askdirectory()
    if directory:
        entry_directory.delete(0, tk.END)
        entry_directory.insert(0, directory)

# Função para executar a unificação
def execute():
    directory = entry_directory.get()
    if not os.path.isdir(directory):
        messagebox.showerror("Erro", "Diretório inválido!")
        return
    file_type = file_type_select.get()  # Obter o tipo de arquivo selecionado
    process_files(directory, check_remove_lines.get(), check_insert_lines.get(), check_add_counter.get(), file_type)

# Configuração da interface
root = tk.Tk()
root.title("Unificação de Arquivos")
root.resizable(False, False)  # Impede redimensionamento da janela
root.geometry("500x500")

# Frame para o diretório
frame_dir = ttk.Frame(root)
frame_dir.pack(pady=10, padx=10, fill=tk.X)

label_dir = ttk.Label(frame_dir, text="Diretório:")
label_dir.pack(side=tk.LEFT)
entry_directory = ttk.Entry(frame_dir, width=50)
entry_directory.pack(side=tk.LEFT, padx=5)
btn_browse = ttk.Button(frame_dir, text="Procurar", command=select_directory)
btn_browse.pack(side=tk.LEFT)

# Adicionando o ComboBox para selecionar tipo de arquivo
file_type_label = ttk.Label(root, text="Selecione o tipo de arquivo:")
file_type_label.pack(pady=5)
file_type_select = ttk.Combobox(root, values=["Texto", "PDF", "Word"], state="readonly", width=20)
file_type_select.set("Texto")  
file_type_select.pack(pady=5)

# Frames de opções de processamento
check_remove_lines = tk.BooleanVar()
check_insert_lines = tk.BooleanVar()
check_add_counter = tk.BooleanVar()

frame_options = ttk.Frame(root)
frame_options.pack(pady=5, padx=10, fill=tk.X)

chk_remove = ttk.Checkbutton(frame_options, text="Remover linhas do início e fim", variable=check_remove_lines)
chk_remove.pack(anchor="w")
chk_insert = ttk.Checkbutton(frame_options, text="Inserir linha no início e fim", variable=check_insert_lines)
chk_insert.pack(anchor="w")
chk_counter = ttk.Checkbutton(frame_options, text="Adicionar contador de linhas no rodapé", variable=check_add_counter)
chk_counter.pack(anchor="w")

# Botão para executar o processo
btn_execute = ttk.Button(root, text="Executar", command=execute)
btn_execute.pack(pady=10)

# log
text_log = scrolledtext.ScrolledText(root, width=60, height=10, font=("Arial", 9))
text_log.pack(padx=10, pady=10)

root.mainloop()
